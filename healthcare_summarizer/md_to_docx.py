from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def markdown_to_docx(markdown_text, output_path):
    """
    Convert Markdown text to DOCX file.
    Handles: headings, lists, tables, bold, italic, code blocks, paragraphs
    
    Args:
        markdown_text: Markdown formatted string
        output_path: Path where DOCX file should be saved
    
    Returns:
        str: Path to created DOCX file
    """
    doc = Document()
    
    # Split into lines for processing
    lines = markdown_text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Check for table
        if is_table_line(line):
            i = process_markdown_table(doc, lines, i)
            continue
        
        # Check for headings
        if line.startswith('#'):
            process_heading(doc, line)
            i += 1
            continue
        
        # Check for unordered list
        if re.match(r'^[\s]*[-*+]\s', line):
            i = process_unordered_list(doc, lines, i)
            continue
        
        # Check for ordered list
        if re.match(r'^[\s]*\d+\.\s', line):
            i = process_ordered_list(doc, lines, i)
            continue
        
        # Check for code block
        if line.strip().startswith('```'):
            i = process_code_block(doc, lines, i)
            continue
        
        # Check for horizontal rule
        if re.match(r'^[-*_]{3,}$', line.strip()):
            doc.add_paragraph('_' * 50)
            i += 1
            continue
        
        # Regular paragraph with inline formatting
        process_paragraph(doc, line)
        i += 1
    
    doc.save(output_path)
    print(f"Markdown converted to DOCX: {output_path}")
    return output_path


def is_table_line(line):
    """Check if line is part of a markdown table."""
    return '|' in line


def process_markdown_table(doc, lines, start_idx):
    """
    Process markdown table and add to document.
    
    Markdown table format:
    | Header 1 | Header 2 |
    |----------|----------|
    | Cell 1   | Cell 2   |
    """
    table_lines = []
    i = start_idx
    
    # Collect all table lines
    while i < len(lines) and '|' in lines[i]:
        table_lines.append(lines[i])
        i += 1
    
    if len(table_lines) < 2:  # Need at least header and separator
        return i
    
    # Parse table
    headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
    
    # Skip separator line (index 1)
    data_rows = []
    for line in table_lines[2:]:
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells:  # Only add non-empty rows
            data_rows.append(cells)
    
    # Create table in document
    if headers and data_rows:
        table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # Add headers
        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = header
            # Make header bold
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add data
        for row_idx, row_data in enumerate(data_rows, start=1):
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(table.rows[row_idx].cells):
                    table.rows[row_idx].cells[col_idx].text = cell_data
        
        doc.add_paragraph()  # Spacing after table
    
    return i


def process_heading(doc, line):
    """Process markdown heading."""
    match = re.match(r'^(#{1,6})\s+(.+)$', line)
    if match:
        level = len(match.group(1))
        text = match.group(2).strip()
        
        # Map markdown levels to Word heading levels
        heading_level = min(level, 3)  # Word supports up to 9, but we'll use 1-3
        heading = doc.add_heading(text, heading_level)
        
        # Center level 1 headings (optional)
        if level == 1:
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER


def process_unordered_list(doc, lines, start_idx):
    """Process markdown unordered list."""
    i = start_idx
    
    while i < len(lines):
        line = lines[i]
        match = re.match(r'^([\s]*)[-*+]\s+(.+)$', line)
        
        if not match:
            break
        
        indent = len(match.group(1))
        text = match.group(2).strip()
        
        # Add as bullet point
        paragraph = doc.add_paragraph(style='List Bullet')
        add_inline_formatting(paragraph, text)
        
        # Handle indentation (nested lists)
        if indent > 0:
            paragraph.paragraph_format.left_indent = Inches(0.5 * (indent // 2))
        
        i += 1
    
    return i


def process_ordered_list(doc, lines, start_idx):
    """Process markdown ordered list."""
    i = start_idx
    
    while i < len(lines):
        line = lines[i]
        match = re.match(r'^([\s]*)\d+\.\s+(.+)$', line)
        
        if not match:
            break
        
        indent = len(match.group(1))
        text = match.group(2).strip()
        
        # Add as numbered list
        paragraph = doc.add_paragraph(style='List Number')
        add_inline_formatting(paragraph, text)
        
        # Handle indentation
        if indent > 0:
            paragraph.paragraph_format.left_indent = Inches(0.5 * (indent // 2))
        
        i += 1
    
    return i


def process_code_block(doc, lines, start_idx):
    """Process markdown code block (```...```)."""
    i = start_idx + 1  # Skip opening ```
    code_lines = []
    
    while i < len(lines) and not lines[i].strip().startswith('```'):
        code_lines.append(lines[i])
        i += 1
    
    if code_lines:
        code_text = '\n'.join(code_lines)
        paragraph = doc.add_paragraph(code_text)
        
        # Style as code
        for run in paragraph.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        
        # Light background (optional - requires more complex styling)
        paragraph.paragraph_format.left_indent = Inches(0.5)
    
    return i + 1  # Skip closing ```


def process_paragraph(doc, line):
    """Process regular paragraph with inline formatting."""
    paragraph = doc.add_paragraph()
    add_inline_formatting(paragraph, line)


def add_inline_formatting(paragraph, text):
    """
    Add inline formatting (bold, italic, code) to paragraph.
    Handles: **bold**, *italic*, `code`, ***bold+italic***
    """
    # Clear any existing runs
    paragraph.text = ''
    
    # Pattern for inline formatting
    # Matches: **bold**, *italic*, `code`, ***bold+italic***
    pattern = r'(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
    
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        
        # Bold and Italic
        if part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        
        # Bold
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        
        # Italic
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        
        # Code
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        
        # Regular text
        else:
            paragraph.add_run(part)


def format_key_as_heading(key):
    """Convert snake_case to Title Case."""
    return key.replace('_', ' ').title()